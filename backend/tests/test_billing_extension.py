"""The paid moment and the extension bridge.

Two things must hold no matter what the client says:
  1. a DOCX never leaves the server unless the bank is unlocked, and
  2. credits appear only via a verified payment, never because a browser asked nicely.
"""
import pytest

from helpers import answer_via_extension, wait_for

# env + the shared `client` fixture live in conftest.py

BANK = "1. Define normalization in DBMS. (5 marks)\n2. Calculate 12 * 4. (2 marks)\n"


@pytest.fixture(scope="module")
def auth(client):
    r = client.post("/api/auth/register", json={
        "email": "buyer@test.dev", "name": "Buyer One", "password": "supersecret1",
    })
    assert r.status_code == 200, r.text
    return {"Authorization": f"Bearer {r.json()['access_token']}"}


def _started_project(client, auth, title, text=BANK):
    """Upload a bank and start it. Questions end up parked for the browser."""
    r = client.post("/api/projects", data={"title": title, "text": text}, headers=auth)
    pid = r.json()["id"]
    p = wait_for(client, auth, pid, lambda p: p["status"] in ("review", "error"))
    assert p["status"] == "review", p
    client.post(f"/api/projects/{pid}/start", headers=auth)
    wait_for(client, auth, pid, lambda p: p["status"] == "done" or p["counts"].get("assist_waiting"))
    return pid


def _answered_project(client, auth, title, text=BANK):
    """...and then let the stand-in extension answer them."""
    pid = _started_project(client, auth, title, text)
    answer_via_extension(client, auth, pid)
    wait_for(client, auth, pid, lambda p: p["status"] == "done")
    return pid


# ---------------- the export paywall ----------------


def test_first_bank_is_free_then_export_costs_a_credit(client, auth):
    first = _answered_project(client, auth, "Free Bank")
    r = client.get(f"/api/projects/{first}/export", headers=auth)
    assert r.status_code == 200
    assert r.content[:2] == b"PK"  # a real .docx is a zip

    # re-download of an already-unlocked bank stays free forever
    assert client.get(f"/api/projects/{first}/export", headers=auth).status_code == 200

    second = _answered_project(client, auth, "Paid Bank")
    r = client.get(f"/api/projects/{second}/export", headers=auth)
    assert r.status_code == 402
    detail = r.json()["detail"]
    assert detail["code"] == "payment_required"
    assert detail["credits"] == 0
    assert detail["packs"], "the paywall must carry the packs it wants to sell"


def test_credits_require_a_completed_order(client, auth):
    balance = client.get("/api/billing/me", headers=auth).json()
    assert balance["credits"] == 0

    order = client.post("/api/billing/checkout", json={"credits": 1}, headers=auth).json()
    assert order["amount_inr"] == 20

    # an order that exists but is unpaid grants nothing
    assert client.get("/api/billing/me", headers=auth).json()["credits"] == 0

    client.get(order["pay_url"])  # mock gateway stands in for the webhook
    assert client.get("/api/billing/me", headers=auth).json()["credits"] == 1

    # replaying the gateway callback must not mint a second credit
    client.get(order["pay_url"])
    assert client.get("/api/billing/me", headers=auth).json()["credits"] == 1


def test_paid_credit_unlocks_the_second_bank(client, auth):
    pid = [p["id"] for p in client.get("/api/projects", headers=auth).json()
           if p["title"] == "Paid Bank"][0]
    assert client.get(f"/api/projects/{pid}/export", headers=auth).status_code == 200
    assert client.get("/api/billing/me", headers=auth).json()["credits"] == 0

    ledger = client.get("/api/billing/history", headers=auth).json()
    assert [t["reason"] for t in ledger][:2] == ["spend", "purchase"]
    assert ledger[0]["balance_after"] == 0


def test_webhook_rejects_an_unsigned_payload(client):
    r = client.post("/api/billing/webhook", json={"event": "payment_link.paid"})
    assert r.status_code == 400


# ---------------- extension bridge ----------------


EXT_BANK = ("1. Explain ACID properties with an example. (10 marks)\n"
            "2. Compute the median of 3, 9, 4, 7. (2 marks)\n")


def test_every_uncached_question_is_parked_for_the_students_own_ai(client, auth):
    pid = _started_project(client, auth, "Extension Bank", text=EXT_BANK)
    p = client.get(f"/api/projects/{pid}", headers=auth).json()
    # the server answers nothing itself — every uncached question waits for the browser
    assert p["counts"].get("assist_waiting") == p["total"]
    assert p["counts"].get("answered") is None

    work = client.get(f"/api/extension/work?project_id={pid}", headers=auth).json()
    assert work["done"] is False
    assert work["prompt"], "the extension must receive a ready-made prompt"
    assert "<question>" in work["prompt"]
    from app.config import get_extension_config
    assert work["preferred_site"] in get_extension_config()["sites"]

    # answering through the normal assist route clears it
    r = client.post(f"/api/questions/{work['question_id']}/assist",
                    json={"content_md": "**Given** 3, 9, 4, 7 → sorted 3, 4, 7, 9\n\nMedian = (4+7)/2\n\nFINAL: 5.5\n\n```verify\n(4+7)/2\n```"}, headers=auth)
    assert r.status_code == 200
    after = client.get(f"/api/extension/work?project_id={pid}", headers=auth).json()
    assert after["question_id"] != work["question_id"]


def test_cache_hits_skip_the_browser(client, auth):
    """A question the class already answered costs nobody anything — and specifically
    costs no trip through the student's browser. The cache outranks everything."""
    pid = _started_project(client, auth, "Cached Bank")
    p = client.get(f"/api/projects/{pid}", headers=auth).json()
    assert p["counts"].get("answered") == p["total"]
    assert p["counts"].get("assist_waiting") is None
    assert all(q["answer"]["engine"] == "cache" for q in p["questions"])


def test_extension_endpoints_require_auth(client):
    assert client.get("/api/extension/work").status_code == 401
    assert client.get("/api/extension/batch").status_code == 401
    assert client.get("/api/extension/config").status_code == 401


# ---------------- batching across assistants ----------------


BIG_BANK = "".join(
    f"{i}. Question number {i} about distributed systems and consensus. ({i} marks)\n"
    for i in range(1, 10)
)


def test_a_batch_follows_the_router_even_when_that_means_one_assistant(client, auth):
    """Three questions at once, each going wherever the router sent it — including all
    three to the same assistant.

    An earlier version forced every batch across three *distinct* sites to spread load.
    That quietly overrode the routing decision on two questions out of every three, which
    trades the thing the router exists for against a rate limit nobody had hit. Spread now
    falls out of the routing itself: different question types go to different sites.
    """
    pid = _started_project(client, auth, "Batch Bank", text=BIG_BANK)

    from app.config import get_extension_config
    theory = get_extension_config()["routing"]["theory"]
    # theory is the bulk type, so its fallback ROTATES across sites with text headroom —
    # a bank keyword-routed while the model's quota is spent must not pile on one site
    theory_sites = set(theory if isinstance(theory, list) else [theory])

    batch = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    assert len(batch) == 3
    # every question in BIG_BANK is theory; each goes where the router (here: the
    # rotating fallback) sent it, and nothing downstream reassigns it
    assert {b["site"] for b in batch} <= theory_sites
    if len(theory_sites) > 1:
        assert len({b["site"] for b in batch}) > 1, "the rotation must actually rotate"
    assert all(b["site"] == b["route_site"] for b in batch), "no slot may override the router"
    assert all(b["prompt"] and "<question>" in b["prompt"] for b in batch)

    # and they are still three separate chats, in paper order
    assert [b["idx"] for b in batch] == [1, 2, 3]

    # leased questions are not handed out twice
    second = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    assert not ({b["question_id"] for b in batch} & {b["question_id"] for b in second})


def test_a_batch_only_uses_assistants_the_student_is_signed_into(client, auth):
    """The one thing that may override the router: an assistant the student can't reach.
    A second-best answer beats no answer."""
    pid = _started_project(client, auth, "Excluded Bank", text=BIG_BANK)
    from app.config import get_extension_config
    everyone = list(get_extension_config()["sites"])
    others = ",".join(s for s in everyone if s != "chatgpt")

    res = client.get(f"/api/extension/batch?project_id={pid}&exclude={others}",
                     headers=auth).json()
    assert [b["site"] for b in res["batch"]] == ["chatgpt", "chatgpt", "chatgpt"]

    res = client.get(f"/api/extension/batch?project_id={pid}&exclude={','.join(everyone)}",
                     headers=auth).json()
    assert res["batch"] == [] and res["error"] == "no_sites_available"


def test_a_dead_tab_does_not_strand_its_question(client, auth):
    """A leased question whose tab was closed must come back to the pool, or the bank
    could never finish."""
    from app.services import queue
    pid = _started_project(client, auth, "Lease Bank", text=BIG_BANK)
    leased = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    assert leased

    p = client.get(f"/api/projects/{pid}", headers=auth).json()
    assert p["counts"].get("assist_running") == len(leased)

    original, queue.LEASE_TTL_S = queue.LEASE_TTL_S, -1   # every lease is now overdue
    try:
        # earlier tests in this module also hold leases, so only the project matters here
        assert queue.expire_leases() >= len(leased)
    finally:
        queue.LEASE_TTL_S = original

    p = client.get(f"/api/projects/{pid}", headers=auth).json()
    assert p["counts"].get("assist_running") is None
    assert p["counts"]["assist_waiting"] == p["total"]      # all back in the pool

    # and they can be handed out again
    again = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    assert len(again) == 3


def test_one_students_bank_is_invisible_to_another(client, auth):
    other = client.post("/api/auth/register", json={
        "email": "other@test.dev", "name": "Other", "password": "supersecret1",
    }).json()
    hdr = {"Authorization": f"Bearer {other['access_token']}"}
    mine = client.get("/api/projects", headers=auth).json()[0]["id"]
    assert client.get(f"/api/projects/{mine}/export", headers=hdr).status_code == 404
    assert client.get(f"/api/extension/work?project_id={mine}", headers=hdr).json()["done"] is True


# ---------------- figures ----------------


def _docx_with_figure():
    """A one-question DOCX carrying an inline image, built in memory."""
    import io

    import docx
    from docx.shared import Inches
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (400, 300), (30, 60, 120)).save(buf, format="PNG")

    d = docx.Document()
    d.add_paragraph("1. Define normalization in DBMS. (5 marks)")
    d.add_paragraph("2. From the graph in Fig. 1, determine the time constant. (10 marks)")
    d.add_picture(io.BytesIO(buf.getvalue()), width=Inches(3))
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def test_a_figure_survives_upload_and_lands_on_the_right_question(client, auth):
    r = client.post("/api/projects", headers=auth,
                    data={"title": "Figure Bank"},
                    files={"file": ("bank.docx", _docx_with_figure(),
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200, r.text
    pid = r.json()["id"]
    p = wait_for(client, auth, pid, lambda p: p["status"] in ("review", "error"))
    assert p["status"] == "review", p

    by_idx = {q["idx"]: q for q in p["questions"]}
    # the picture follows Q2 in the document, so it belongs to Q2 — not Q1
    assert by_idx[0]["figures"] == []
    assert len(by_idx[1]["figures"]) == 1

    # and a question that names a figure it HAS is not flagged as missing one
    assert by_idx[1]["needs_figure"] is False

    # the image is actually served
    img = client.get(by_idx[1]["figures"][0]["url"], headers=auth)
    assert img.status_code == 200
    assert img.content[:4] in (b"\x89PNG", b"\xff\xd8\xff\xe0")


def test_a_question_naming_a_figure_it_lacks_is_flagged(client, auth):
    bank = ("1. Explain the OSI model layers. (10 marks)\n"
            "2. From the graph shown above, find the peak voltage. (5 marks)\n")
    r = client.post("/api/projects", data={"title": "No Figure Bank", "text": bank}, headers=auth)
    pid = r.json()["id"]
    p = wait_for(client, auth, pid, lambda p: p["status"] == "review")
    by_idx = {q["idx"]: q for q in p["questions"]}
    assert by_idx[0]["needs_figure"] is False
    assert by_idx[1]["needs_figure"] is True, "a figure reference with no figure must be flagged"


def test_an_uploaded_paper_is_answered_from_the_document_itself(client, auth):
    """Every question in a bank we still hold the file for is answered against that file,
    not against our extracted text — because the file is simply a better copy of the
    question. It carries the figures, the tables and the original wording, and the model
    reading it decides which picture belongs to which question far better than we can.

    Extraction stays responsible for how many questions there are; never for what they say.
    """
    pid = [p["id"] for p in client.get("/api/projects", headers=auth).json()
           if p["title"] == "Figure Bank"][0]
    client.post(f"/api/projects/{pid}/start", headers=auth)
    wait_for(client, auth, pid, lambda p: p["counts"].get("assist_waiting"))

    batch = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    assert batch
    assert all(b.get("document") for b in batch), "a file-backed bank answers from the file"

    item = [b for b in batch if b["document"]["number"] == 2][0]   # it was Q2 in the file
    assert item["document"]["url"].endswith(pid)
    assert "<answer_question>2</answer_question>" in item["prompt"]
    assert "Answer ONLY that one question" in item["prompt"]
    # a paper the AI can't find the question in must not cost us the question
    assert "<question>" in item["fallback_prompt"]

    # and the file itself is downloadable by its owner
    doc = client.get(item["document"]["url"], headers=auth)
    assert doc.status_code == 200
    assert doc.content[:2] == b"PK"                 # the docx we uploaded


def test_the_review_save_keeps_figures_and_the_number_in_the_paper(client, auth):
    """Saving the review must not cost a question the two things the student cannot retype.

    This is the regression that quietly killed every image question: the save replaced the
    rows, which set each figure's question_id to NULL and dropped the number the question
    carried in the file. Both are invisible on the review screen, so nothing looked wrong —
    the AI simply answered figure questions without ever seeing the figure.
    """
    r = client.post("/api/projects", headers=auth,
                    data={"title": "Review Keeps Figures"},
                    files={"file": ("bank.docx", _docx_with_figure(),
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    pid = r.json()["id"]
    p = wait_for(client, auth, pid, lambda p: p["status"] == "review")

    figure_q = [q for q in p["questions"] if q["figures"]][0]
    assert figure_q["source_number"] == 2, "the review screen needs the file's own numbering"

    # a normal review: fix a typo on one question, leave the rest alone
    saved = client.put(f"/api/projects/{pid}/questions", headers=auth, json={"questions": [
        {"id": q["id"], "text": q["text"] + " Explain briefly.", "marks": q["marks"],
         "number": q["source_number"]}
        for q in p["questions"]
    ]}).json()

    after = [q for q in saved["questions"] if q["id"] == figure_q["id"]][0]
    assert after["figures"], "the figure must still belong to its question after review"
    assert after["source_number"] == 2

    # and a question the student deletes really goes
    keep = [q for q in saved["questions"] if q["id"] != figure_q["id"]]
    trimmed = client.put(f"/api/projects/{pid}/questions", headers=auth, json={
        "questions": [{"id": q["id"], "text": q["text"], "marks": q["marks"]} for q in keep]}).json()
    assert len(trimmed["questions"]) == len(keep)
    assert figure_q["id"] not in [q["id"] for q in trimmed["questions"]]


def test_a_cant_answer_reply_is_never_accepted_as_an_answer(client, auth):
    """The incident this guards against: a document upload failed, the AI replied "No file
    appears to be attached to this chat", and that sentence was accepted, stored, CACHED,
    and served 15 more times. A refusal must bounce — the question goes back to waiting
    and the extension retries elsewhere."""
    pid = _started_project(client, auth, "Refusal Bank",
                           text="1. Explain the OSI model layers in detail. (10 marks)\n"
                                "2. Compare TCP and UDP protocols. (5 marks)\n"
                                "3. Define normalization with examples. (5 marks)\n")
    work = client.get(f"/api/extension/work?project_id={pid}", headers=auth).json()

    for bad in [
        "No file appears to be attached to this chat — the uploads folder is empty. Please re-upload it.",
        "I cannot see the document you're referring to. Could you please attach the file again?",
        "The question content (image or text) is missing from the provided prompt. To proceed, please paste it.",
        "NOT_FOUND",
        "Sure, happy to help!",   # shrug-length is not an answer either
    ]:
        r = client.post(f"/api/questions/{work['question_id']}/assist",
                        json={"content_md": bad}, headers=auth)
        assert r.status_code == 422, f"accepted a non-answer: {bad!r}"

    # the question is back in the pool, not stuck answered-with-garbage
    p = client.get(f"/api/projects/{pid}", headers=auth).json()
    q = [x for x in p["questions"] if x["id"] == work["question_id"]][0]
    assert q["status"] == "assist_waiting"

    # and a real answer that merely DISCUSSES missing data sails through
    real = ("**Definition:** normalization organises relational tables to remove redundancy.\n\n"
            "If an attribute's value is missing from the provided data, 1NF still requires the "
            "cell to hold an atomic value.\n\n**Takeaway:** normalize to 3NF for exam answers.")
    r = client.post(f"/api/questions/{work['question_id']}/assist",
                    json={"content_md": real}, headers=auth)
    assert r.status_code == 200


def test_refusals_and_shrugs_never_enter_the_class_cache(client, auth):
    """Belt to the endpoint's braces: even if a refusal reaches cache.store by another
    path, it must not be stored — a cached wrong answer is wrong for the whole class."""
    from app.db import SessionLocal
    from app.services import cache

    db = SessionLocal()
    try:
        for bad in ["No file appears to be attached to this chat, please re-upload.",
                    "Too short."]:
            cache.store(db, "unique refusal-cache probe q?", 5, "theory",
                        content_md=bad, provider="assist", model="x", verified=None)
            assert cache.lookup(db, "unique refusal-cache probe q?", 5, "theory") is None
    finally:
        db.close()


def test_a_whole_bank_is_routed_in_one_call(client, auth):
    """Routing costs one call per *bank*, not one per question.

    A free tier's daily allowance is measured in tens, so a 28-question bank routed one
    question at a time runs out of routing before it runs out of questions — and the
    fallback is keywords, which is exactly the judgement we were paying for."""
    from app.services import providers, router_agent

    calls = []
    original = providers.chat

    async def counting(provider, model, messages, **kw):
        calls.append(messages[0]["content"][:20])
        return await original(provider, model, messages, **kw)

    providers.chat = counting
    try:
        import asyncio
        routes = asyncio.run(
            router_agent.classify_many([f"{i}. Explain concept {i}." for i in range(12)]))
    finally:
        providers.chat = original

    assert len(routes) == 12
    assert all(r["qtype"] in router_agent.QTYPES for r in routes)
    assert all(r["site"] for r in routes)
    assert len(calls) == 1, f"12 questions must cost one call, not {len(calls)}"


def test_a_question_needing_the_paper_goes_where_uploads_are_free(client, auth):
    """Document mode uploads the question paper once PER QUESTION, and a free ChatGPT
    account allows three uploads a day — spent before the first batch of three finishes.
    So a question that needs the paper is moved to a site with upload headroom, even when
    the router preferred someone else for the content.

    This is the constraint that silently kills a 28-question run: everything looks routed
    correctly right up to the point the uploads stop being accepted."""
    from app.config import get_extension_config

    cfg = get_extension_config()
    roomy = {k for k, v in cfg["sites"].items() if v.get("documents") in ("unlimited", "generous")}
    assert roomy, "at least one site must be able to take a document at volume"
    assert cfg["document_sites"][0] in roomy, "the first choice must have headroom"
    assert cfg["document_sites"][-1] not in roomy, "the tightest site belongs last"

    r = client.post("/api/projects", headers=auth,
                    data={"title": "Upload Budget Bank"},
                    files={"file": ("bank.docx", _docx_with_figure(),
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    pid = r.json()["id"]
    wait_for(client, auth, pid, lambda p: p["status"] == "review")
    client.post(f"/api/projects/{pid}/start", headers=auth)
    wait_for(client, auth, pid, lambda p: p["counts"].get("assist_waiting"))

    batch = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    doc_items = [b for b in batch if b.get("document")]
    assert doc_items, "a bank we still hold the file for is answered from that file"
    for b in doc_items:
        assert b["site"] in roomy, f"Q{b['idx']} would burn a scarce upload on {b['site']}"


def test_figure_questions_are_grouped_apart_from_theory(client, auth):
    """The deck shows figure questions in their own row. They're the ones worth a second
    look: a misread diagram becomes a confident wrong answer rather than an obvious blank,
    so burying them among thirty theory questions is exactly where they get missed."""
    pid = _started_project(
        client, auth, "Mixed Bank",
        text="1. Define normalization in DBMS. (5 marks)\n"
             "2. From the graph shown above, find the peak voltage. (5 marks)\n"
             "3. Compare TCP and UDP. (5 marks)\n"
             "4. Draw the ER diagram for a library system. (10 marks)\n")
    qs = client.get(f"/api/projects/{pid}", headers=auth).json()["questions"]
    visual = {q["idx"]: q["visual"] for q in qs}

    assert visual[1] is True, "a question reading a graph belongs in the figure row"
    assert visual[3] is True, "so does one that has to draw a diagram"
    assert visual[0] is False and visual[2] is False


def test_a_repeated_number_is_not_trusted_to_identify_a_question():
    """Spreadsheet exports and multi-section papers restart their numbering, so a paper can
    hold two question 11s. "Answer question 11" is then a coin toss — the number has to be
    thrown away and the question quoted instead."""
    from types import SimpleNamespace

    from app.services.paper import number_is_unique

    def q(number):
        return SimpleNamespace(source_number=number, project=project)

    project = SimpleNamespace()
    a, b, c = q(11), q(11), q(13)
    project.questions = [a, b, c]

    assert number_is_unique(c) is True
    assert number_is_unique(a) is False
    assert number_is_unique(b) is False
    assert number_is_unique(q(None)) is False


def test_an_unnumbered_question_is_located_by_quoting_it(client, auth):
    """Papers without usable numbering — a photo of a sheet, a bulleted list — still get
    document mode. We quote the question's opening instead of naming a number, so their
    figure questions aren't the one case that falls through."""
    from app.services import solver

    prompt = solver.build_document_prompt(
        "From the waveform shown, determine the peak-to-peak voltage.", "numerical", 5)
    assert "<answer_question>From the waveform shown" in prompt
    assert "Answer ONLY that one question" in prompt
    assert "question None" not in prompt


def test_the_document_is_not_offered_when_there_is_no_source_file(client, auth):
    """A bank pasted as text has no file to attach, so those questions must fall back to
    the ordinary prompt rather than referencing a document that doesn't exist."""
    pid = _started_project(client, auth, "Pasted With Figure Ref",
                           text="1. From the graph above, find the peak voltage. (5 marks)\n"
                                "2. Define normalization. (5 marks)\n"
                                "3. Compare TCP and UDP. (5 marks)\n")
    batch = client.get(f"/api/extension/batch?project_id={pid}", headers=auth).json()["batch"]
    assert batch
    assert all(not b.get("document") for b in batch)
    assert all("<question>" in b["prompt"] for b in batch)


def test_a_figure_is_never_leaked_across_accounts(client, auth):
    pid = [p["id"] for p in client.get("/api/projects", headers=auth).json()
           if p["title"] == "Figure Bank"][0]
    fid = [q for q in client.get(f"/api/projects/{pid}", headers=auth).json()["questions"]
           if q["figures"]][0]["figures"][0]["id"]
    other = client.post("/api/auth/register", json={
        "email": "peeker@test.dev", "name": "Peeker", "password": "supersecret1"}).json()
    hdr = {"Authorization": f"Bearer {other['access_token']}"}
    assert client.get(f"/api/figures/{fid}", headers=hdr).status_code == 404
    assert client.get(f"/api/figures/{fid}").status_code == 401
