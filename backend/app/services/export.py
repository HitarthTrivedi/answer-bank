"""Answers → one polished .docx. A small markdown renderer that understands exactly
the dialect our solver contract produces: headings, bold/italic/inline-code, lists,
tables, code fences, $$ math (rendered via mathtext), graphspec fences (rendered
server-side) and mermaid fences (client-posted PNG assets, else code fallback)."""
import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..models import AnswerAsset, Project
from . import diagrams

_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\$[^$\n]+\$)")


def _add_inline(paragraph, text: str) -> None:
    """Markdown inline formatting → docx runs. Inline math falls back to plain text."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("$") and part.endswith("$") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(code.rstrip("\n"))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_image(doc: Document, png: bytes, width_in: float = 5.5) -> None:
    doc.add_picture(io.BytesIO(png), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


_FENCE_OPEN = re.compile(r"^```(\w+)?\s*$")


def _render_markdown(doc: Document, md: str, assets_by_key: dict[str, str]) -> None:
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        fence = _FENCE_OPEN.match(line)

        if fence:  # collect the fenced block
            lang = (fence.group(1) or "").lower()
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # closing fence
            body = "\n".join(block)
            if lang == "verify":
                continue  # machine-check plumbing, not for the reader
            if lang == "graphspec":
                png = diagrams.render_graphspec(body)
                if png:
                    _add_image(doc, png)
                continue
            if lang == "mermaid":
                asset_path = assets_by_key.get(diagrams.spec_key(body))
                if asset_path:
                    try:
                        with open(asset_path, "rb") as f:
                            _add_image(doc, f.read())
                        continue
                    except OSError:
                        pass
                _add_code_block(doc, body)  # fallback: show the mermaid source
                continue
            _add_code_block(doc, body)
            continue

        if line.strip().startswith("$$"):  # display math (single or multi-line)
            math_lines = [line.strip().strip("$")]
            if not (line.strip().endswith("$$") and len(line.strip()) > 2):
                i += 1
                while i < len(lines) and "$$" not in lines[i]:
                    math_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    math_lines.append(lines[i].replace("$$", ""))
            latex = " ".join(m for m in math_lines if m.strip())
            png = diagrams.render_mathtext(latex)
            if png:
                _add_image(doc, png, width_in=min(4.5, 0.12 * max(10, len(latex))))
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(latex).italic = True
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):  # markdown table
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):  # skip separator row
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        _add_inline(table.rows[r].cells[c].paragraphs[0], cell)
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            _add_inline(doc.add_heading("", level=min(level + 1, 4)), stripped.lstrip("#").strip())
        elif re.match(r"^[-*]\s+", stripped):
            _add_inline(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+[\.\)]\s+", stripped):
            _add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+[\.\)]\s+", "", stripped))
        elif stripped.upper().startswith("FINAL:"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        else:
            _add_inline(doc.add_paragraph(), stripped)
        i += 1


def build_docx(project: Project, db, buyer: str = "") -> bytes:
    doc = Document()

    # cover
    title = doc.add_heading(project.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answered = [q for q in project.questions if q.answer is not None]
    sub = doc.add_paragraph(
        f"Answer document · {len(answered)} of {len(project.questions)} questions answered\n"
        f"Generated by AnswerBank · {datetime.now(timezone.utc).strftime('%d %b %Y')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if buyer:
        # names the copy. Not DRM — just enough friction that forwarding it round the
        # class feels like handing over something with your name on it.
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run(f"Prepared for {buyer}")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_page_break()

    # index
    doc.add_heading("Questions", level=1)
    for q in project.questions:
        p = doc.add_paragraph(style="List Number")
        p.add_run(q.text[:120] + ("…" if len(q.text) > 120 else ""))
        if q.marks:
            p.add_run(f"  [{q.marks} marks]").italic = True
    doc.add_page_break()

    # answers
    for q in project.questions:
        head = doc.add_heading("", level=1)
        head.add_run(f"Q{q.idx + 1}. ")
        _add_inline(head, q.text)
        if q.marks:
            meta = doc.add_paragraph()
            meta.add_run(f"[{q.marks} marks]").italic = True

        if q.answer is None:
            doc.add_paragraph("— not answered yet —").runs[0].italic = True
        else:
            assets = db.query(AnswerAsset).filter_by(answer_id=q.answer.id).all()
            assets_by_key = {a.key: a.path for a in assets}
            _render_markdown(doc, q.answer.content_md, assets_by_key)
            credit = doc.add_paragraph()
            label = {"api": q.answer.model, "assist": "student's own AI (assist mode)",
                     "cache": "class cache"}.get(q.answer.engine, q.answer.engine)
            badge = " · verified ✓" if q.answer.verified else (" · check working ⚠" if q.answer.verified is False else "")
            run = credit.add_run(f"— answered via {label}{badge}")
            run.font.size = Pt(8)
            run.italic = True
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
