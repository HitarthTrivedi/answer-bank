"""File → raw text + the figures embedded in it.

Figures matter because a question bank full of "from the graph in Fig. 3" is otherwise
answered blind — the AI never sees the graph and invents a confident wrong answer, which
is the worst possible failure for an exam document.

We do NOT interpret figures here. No OCR, no vision model, no cost. We keep the pixels
and their position, and the student's own browser AI reads them when it answers — it is
better at it than tesseract and it is already paid for.
"""
import io
import logging
import zipfile
from pathlib import Path

from fastapi import HTTPException

from ..config import get_settings

log = logging.getLogger("prism.ingest")

# What separates a figure from a bullet glyph, a rule or a logo is its SIZE ON THE PAGE,
# not its file size. Line-art — exactly what exam diagrams are — compresses to almost
# nothing, so a byte threshold would throw away the clearest circuit diagram while
# keeping a noisy JPEG watermark. Bytes are only a floor against 1x1 junk.
MIN_FIGURE_BYTES = 200
MIN_FIGURE_PX = 120

# A line that is nothing but a number: a table row label, which names the figure sitting
# in that row. Worth this much of a head start over raw distance when it sits below an
# image — see _anchor_figures.
_ROW_LABEL = __import__("re").compile(r"\d{1,3}")
_ROW_LABEL_BONUS = 60.0

MAGIC = {
    "pdf": b"%PDF",
    "docx": b"PK\x03\x04",
    "png": b"\x89PNG",
    "jpg": b"\xff\xd8\xff",
}


def sniff_kind(data: bytes, claimed_ext: str) -> str:
    """Return the real file kind from content, cross-checked against the claimed extension."""
    ext = claimed_ext.lower().lstrip(".")
    if ext == "jpeg":
        ext = "jpg"
    if ext in ("txt", "md"):
        try:
            data.decode("utf-8")
            return "txt"
        except UnicodeDecodeError:
            raise HTTPException(400, "Text file is not valid UTF-8")
    magic = MAGIC.get(ext)
    if magic is None:
        raise HTTPException(400, f"Unsupported file type: .{ext}")
    if not data.startswith(magic):
        raise HTTPException(400, f"File content does not match its .{ext} extension")
    if ext == "docx":
        # a real docx is a zip containing [Content_Types].xml
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                if "[Content_Types].xml" not in z.namelist():
                    raise HTTPException(400, "Not a valid .docx archive")
        except zipfile.BadZipFile:
            raise HTTPException(400, "Not a valid .docx archive")
    return ext


def validate_upload(data: bytes, filename: str) -> str:
    s = get_settings()
    if len(data) > s.max_upload_mb * 1024 * 1024:
        raise HTTPException(413, f"File exceeds {s.max_upload_mb} MB limit")
    if len(data) == 0:
        raise HTTPException(400, "Empty file")
    ext = Path(filename).suffix.lstrip(".").lower()
    if ext not in {e.strip() for e in s.allowed_extensions.split(",")}:
        raise HTTPException(400, f"Extension .{ext} not allowed")
    return sniff_kind(data, ext)


def extract_text(data: bytes, kind: str) -> str:
    return extract_document(data, kind)["text"]


def extract_document(data: bytes, kind: str) -> dict:
    """Returns {text, figures}. `figures` is a list of
    {bytes, ext, anchor} where `anchor` is the character offset in `text` that the figure
    sits nearest to — that offset is what later maps a figure onto a question."""
    if kind == "txt":
        return {"text": data.decode("utf-8"), "figures": []}
    if kind == "pdf":
        return _pdf_document(data)
    if kind == "docx":
        return _docx_document(data)
    if kind in ("png", "jpg"):
        # the upload IS the figure: no OCR, hand the whole thing to the browser AI
        return {"text": _image_text(data), "figures": [{"bytes": data, "ext": kind, "anchor": 0}]}
    raise HTTPException(400, f"Cannot extract text from kind '{kind}'")


def _usable_figure(raw: bytes) -> tuple[bool, str]:
    """Big enough to be a real figure? Returns (ok, extension)."""
    if len(raw) < MIN_FIGURE_BYTES:
        return False, ""
    try:
        from PIL import Image

        with Image.open(io.BytesIO(raw)) as im:
            # both dimensions must be reasonable AND the area substantial, so a thin
            # decorative rule (900x8) is rejected while a small diagram is kept
            if min(im.size) < MIN_FIGURE_PX or im.width * im.height < MIN_FIGURE_PX ** 2 * 2:
                return False, ""
            return True, "png" if im.format == "PNG" else "jpg"
    except Exception:
        return False, ""


def _mul(m: list, n: list) -> list:
    """Compose two PDF matrices [a b c d e f] (m applied first, then n)."""
    a1, b1, c1, d1, e1, f1 = m
    a2, b2, c2, d2, e2, f2 = n
    return [a1 * a2 + b1 * c2, a1 * b2 + b1 * d2,
            c1 * a2 + d1 * c2, c1 * b2 + d1 * d2,
            e1 * a2 + f1 * c2 + e2, e1 * b2 + f1 * d2 + f2]


def _image_positions(page, reader) -> dict[str, tuple[float, float]]:
    """Each image's vertical EXTENT on the page (low, high), by XObject name.

    PDFs don't store "this image belongs to question 7" — they store a transformation
    matrix and a draw call. We replay the content stream, tracking the CTM through the
    q/Q stack, and read the placement at each `Do`.

    The extent matters, not the midpoint. Exam figures are tall, and the row label that
    identifies them ("7") sits somewhere inside that height — often nowhere near the
    centre. Matching on the midpoint sends two images on a sparse page to the same row.
    """
    from pypdf.generic import ContentStream

    out: dict[str, tuple[float, float]] = {}
    ctm = [1.0, 0, 0, 1.0, 0, 0]
    stack: list[list] = []
    for operands, op in ContentStream(page.get_contents(), reader).operations:
        if op == b"q":
            stack.append(list(ctm))
        elif op == b"Q":
            ctm = stack.pop() if stack else [1.0, 0, 0, 1.0, 0, 0]
        elif op == b"cm" and len(operands) == 6:
            ctm = _mul([float(x) for x in operands], ctm)
        elif op == b"Do" and operands:
            name = str(operands[0]).lstrip("/")
            # the image fills the unit square under the CTM; d may be negative (flip)
            bottom, top = ctm[5], ctm[5] + ctm[3]
            out.setdefault(name, (min(bottom, top), max(bottom, top)))
    return out


def _text_positions(page) -> list[tuple[float, str]]:
    """(vertical position, text) for every run of text drawn on the page."""
    spans: list[tuple[float, str]] = []

    def visit(text, cm, tm, font_dict, font_size):
        if text and text.strip():
            spans.append((_mul(list(tm), list(cm))[5], text.strip()))

    try:
        page.extract_text(visitor_text=visit)
    except Exception:
        return []
    return spans


def _anchor_figures(page, reader, body: str, offset: int) -> list[dict]:
    """Work out which line of text each image on this page belongs beside.

    A figure is identified by the nearest text — usually its row label or the question it
    illustrates — but "nearest" alone is not enough. In a table, a label sits a couple of
    points BELOW its own image and only slightly further from the image of the row above,
    so pure proximity hands the same label to both.

    The fix is that a label identifies ONE figure: pair them up closest-first, consuming
    each label as it is used. An image left without a label (two figures in one row) falls
    back to its nearest text, sharing.
    """
    positions = _image_positions(page, reader)
    spans = _text_positions(page)
    images = []
    for img in page.images:
        ok, ext = _usable_figure(img.data)
        if not ok:
            continue
        key = str(img.name).lstrip("/")
        images.append((img, ext, positions.get(key.rsplit(".", 1)[0]) or positions.get(key)))

    if not images:
        return []
    if not spans:
        return [{"bytes": im.data, "ext": ext, "anchor": offset} for im, ext, _ in images]

    def gap(band, y):
        if band is None:
            return float("inf")
        low, high = band
        return 0.0 if low <= y <= high else min(abs(low - y), abs(high - y))

    def score(band, j):
        """Distance, with a thumb on the scale for a table row label.

        Layout cuts both ways: in an exam paper the question OWNING a figure sits above
        it, while in a spreadsheet the row label sits just below its own image. Those are
        opposite directions, so distance alone cannot settle it.

        What separates them is what the text says. A bare number under an image is a row
        label and names that figure. Anything else below a figure is usually the NEXT
        question, so for those we let plain proximity pick the question above.
        """
        y, text = spans[j]
        if band is not None and _ROW_LABEL.fullmatch(text) and y < band[0]:
            return gap(band, y) - _ROW_LABEL_BONUS
        return gap(band, y)

    pairs = sorted(
        ((score(band, j), i, j) for i, (_, _, band) in enumerate(images)
         for j in range(len(spans))),
        key=lambda t: t[0],
    )
    taken_img, taken_span, chosen = set(), set(), {}
    for _, i, j in pairs:
        if i in taken_img or j in taken_span:
            continue
        chosen[i] = j
        taken_img.add(i)
        taken_span.add(j)

    out = []
    for i, (img, ext, band) in enumerate(images):
        anchor = offset
        j = chosen.get(i)
        if j is None and band is not None:          # more figures than labels — share
            j = min(range(len(spans)), key=lambda k: gap(band, spans[k][0]))
        if j is not None:
            where = body.find(spans[j][1])
            if where >= 0:
                anchor = offset + where
        out.append({"bytes": img.data, "ext": ext, "anchor": anchor})
    return out


def _pdf_document(data: bytes) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages, figures, offset = [], [], 0
    for page in reader.pages:
        body = page.extract_text() or ""
        # Anchoring to the START of the page — which is what this used to do — collapses
        # every figure on a page onto one point, so a spreadsheet-style bank with ten rows
        # per page sends all its diagrams to one question and starves the rest.
        try:
            figures.extend(_anchor_figures(page, reader, body, offset))
        except Exception as e:      # a malformed XObject must never fail the upload
            log.warning("could not read images on a page: %s", e)
        pages.append(body)
        offset += len(body) + 2     # the "\n\n" join below

    text = "\n\n".join(pages).strip()
    if len(text) < 40 * max(1, len(pages)) and len(text) < 200:
        raise HTTPException(
            422,
            "This PDF looks scanned (no text layer). Export it with OCR first, "
            "or paste the questions as text.",
        )
    return {"text": text, "figures": figures}


_BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"


def _docx_document(data: bytes) -> dict:
    import docx

    d = docx.Document(io.BytesIO(data))
    lines, figures, offset = [], [], 0
    rels = d.part.related_parts

    for para in d.paragraphs:
        # a paragraph can carry inline images; anchor them where the paragraph sits
        for blip in para._p.iter(_BLIP):
            rid = blip.get(_EMBED)
            part = rels.get(rid) if rid else None
            if part is None:
                continue
            ok, ext = _usable_figure(part.blob)
            if ok:
                figures.append({"bytes": part.blob, "ext": ext, "anchor": offset})
        if para.text.strip():
            lines.append(para.text)
            offset += len(para.text) + 1

    # tables become pipe-delimited text — the house style keeps tabular data AS a table
    for table in d.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
                offset += len(row_text) + 1

    return {"text": "\n".join(lines), "figures": figures}


def _image_text(data: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image

        text = pytesseract.image_to_string(Image.open(io.BytesIO(data)))
        if not text.strip():
            raise HTTPException(422, "OCR found no text in this image")
        return text
    except ImportError:
        raise HTTPException(
            422,
            "Image OCR needs tesseract installed (brew install tesseract && pip install pytesseract). "
            "Until then, paste the questions as text.",
        )
